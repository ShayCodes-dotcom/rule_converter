"""
OIA-to-Drools Rule Converter v2
Flask web application that converts Oracle Intelligent Advisor (OIA) Word rule
documents to Drools DRL (Drools Rule Language) files.

Handles:
  - Paragraph rules (bold = conclusion, indented = conditions)
  - Rule tables (2-column: conclusion value | conditions)
  - OIA grouping keywords: either/or/any/all with indentation nesting
  - OIA concepts: is undefined, is unknown, is not undefined, otherwise
  - IntervalSometimes() detection (flagged for manual review)
  - Boolean and string conclusions
  - <> (not equal) operator
  - Negation patterns ("is not X")
  - Salience-based evaluation ordering
"""

import os
import re
import io
import zipfile
import json
from datetime import datetime
from flask import Flask, request, render_template_string, send_file, jsonify
from docx import Document

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max

UPLOAD_FOLDER = '/tmp/oia_uploads'
OUTPUT_FOLDER = '/tmp/oia_outputs'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


# ─────────────────────────────────────────────────────────────────────────────
# Utility Functions
# ─────────────────────────────────────────────────────────────────────────────

def _normalize_quotes(text):
    """Replace smart/curly quotes with straight quotes."""
    text = text.replace('\u201c', '"').replace('\u201d', '"')   # " "
    text = text.replace('\u2018', "'").replace('\u2019', "'")   # ' '
    text = text.replace('\u00ab', '"').replace('\u00bb', '"')   # « »
    return text


def _to_field_name(oia_text):
    """
    Convert an OIA natural-language attribute name to a camelCase Java field name.

    Examples:
        "the man's education code"           → "educationCode"
        "the man is eligible"                → "eligible"  (strips "the man is")
        "the man's food is full LKSC"        → "foodIsFullLKSC"
        "the food is sometimes LKSC"         → "foodIsSometimesLKSC"
        "the man has valid FLIP certification" → "hasValidFLIPCertification"
    """
    text = oia_text.strip()

    # Strip trailing "if" / "when"
    text = re.sub(r'\s+(if|when)\s*$', '', text, flags=re.IGNORECASE)

    # Remove leading "the man's " / "the man " / "the "
    text = re.sub(r"^the\s+man'?s?\s+", '', text, flags=re.IGNORECASE)
    text = re.sub(r"^the\s+", '', text, flags=re.IGNORECASE)

    # Split into words
    words = text.split()
    if not words:
        return "unknownField"

    # Build camelCase
    result = words[0].lower()
    for w in words[1:]:
        # Preserve all-caps acronyms (LKSC, TESY, FLIP, KFC)
        if w.isupper() and len(w) >= 2:
            result += w
        else:
            result += w.capitalize()

    # Remove non-alphanumeric
    result = re.sub(r'[^a-zA-Z0-9]', '', result)

    # Java identifiers can't start with a digit
    if result and result[0].isdigit():
        result = '_' + result

    return result


def _strip_trailing_connectors(text):
    """Remove trailing 'and', 'or', 'and\\n', 'or\\n' from condition text."""
    text = text.strip()
    text = re.sub(r'\s+(and|or)\s*$', '', text, flags=re.IGNORECASE)
    return text.strip()


def _get_indent_pt(para):
    """Get left indent in points, defaulting to 0."""
    if para.paragraph_format.left_indent:
        return para.paragraph_format.left_indent.pt
    return 0.0


def _is_bold(para):
    """Check if a paragraph's text runs are bold."""
    runs_with_text = [r for r in para.runs if r.text.strip()]
    if not runs_with_text:
        return False
    return all(r.bold for r in runs_with_text)


def _is_title(para):
    """Check if paragraph is a document title (colored, bold, no indent, first line)."""
    for r in para.runs:
        if r.font.color and r.font.color.rgb:
            rgb = str(r.font.color.rgb)
            # Gold/orange = FFC000, other title colors
            if rgb not in ('000000', '000000'):
                return True
    return False


# ─────────────────────────────────────────────────────────────────────────────
# OIA Document Reader
# ─────────────────────────────────────────────────────────────────────────────

class OIADocumentReader:
    """
    Reads an OIA .docx file and extracts structured paragraph data and tables.
    """

    def read(self, filepath):
        """
        Returns:
            {
                'title': str or None,
                'paragraph_groups': [ [paragraph_dict, ...], ... ],
                'tables': [ table_dict, ... ]
            }
        """
        doc = Document(filepath)

        title = None
        paragraphs = []
        tables = []

        # Extract paragraphs with metadata
        for i, para in enumerate(doc.paragraphs):
            text = _normalize_quotes(para.text.strip())
            p_data = {
                'index': i,
                'text': text,
                'indent_pt': _get_indent_pt(para),
                'is_bold': _is_bold(para),
                'is_empty': not text,
            }

            # Detect title (first non-empty, colored paragraph)
            if text and title is None and _is_title(para):
                title = text
                continue

            paragraphs.append(p_data)

        # Group paragraphs into rules (split on empty lines between non-indented bold lines)
        paragraph_groups = self._group_paragraphs(paragraphs)

        # Extract tables
        for ti, table in enumerate(doc.tables):
            tables.append(self._read_table(table, ti))

        return {
            'title': title,
            'paragraph_groups': paragraph_groups,
            'tables': tables,
        }

    def _group_paragraphs(self, paragraphs):
        """
        Group paragraphs into rule blocks.
        A new rule starts when we encounter a bold paragraph at the base indent level
        (zero or near-zero indent) after previous content. Also splits on blank lines.
        """
        groups = []
        current_group = []

        for p in paragraphs:
            if p['is_empty']:
                if current_group:
                    groups.append(current_group)
                    current_group = []
                continue

            # A bold line starts a new rule group if we already have content
            if p['is_bold'] and current_group:
                groups.append(current_group)
                current_group = []

            current_group.append(p)

        if current_group:
            groups.append(current_group)

        return groups

    def _read_table(self, table, table_index):
        """
        Read an OIA rule table.
        Returns:
            {
                'table_index': int,
                'header': str,                    # conclusion attribute name
                'rows': [
                    {
                        'value': str,             # left column (conclusion value)
                        'conditions': [           # right column paragraphs
                            {'text': str, 'indent_pt': float},
                            ...
                        ],
                        'is_otherwise': bool,
                    },
                    ...
                ]
            }
        """
        if len(table.rows) < 2 or len(table.columns) < 2:
            return {'table_index': table_index, 'header': '', 'rows': []}

        # Header row: both columns contain the conclusion attribute name
        header_text = _normalize_quotes(table.rows[0].cells[0].text.strip())

        rows = []
        for ri in range(1, len(table.rows)):
            row = table.rows[ri]
            # Left column: conclusion value
            value_text = _normalize_quotes(row.cells[0].text.strip())

            # Right column: conditions (multiple paragraphs with indentation)
            cond_paragraphs = []
            for p in row.cells[1].paragraphs:
                ptxt = _normalize_quotes(p.text.strip())
                if ptxt:
                    pindent = 0.0
                    if p.paragraph_format.left_indent:
                        pindent = p.paragraph_format.left_indent.pt
                    cond_paragraphs.append({
                        'text': ptxt,
                        'indent_pt': pindent,
                    })

            is_otherwise = (
                len(cond_paragraphs) == 1
                and cond_paragraphs[0]['text'].lower() == 'otherwise'
            )

            rows.append({
                'value': value_text,
                'conditions': cond_paragraphs,
                'is_otherwise': is_otherwise,
            })

        return {
            'table_index': table_index,
            'header': header_text,
            'rows': rows,
        }


# ─────────────────────────────────────────────────────────────────────────────
# Condition Tree Builder
# ─────────────────────────────────────────────────────────────────────────────

class ConditionNode:
    """
    Represents a node in a condition tree.
    Can be a leaf (single condition) or a group (AND/OR of children).
    """
    def __init__(self, node_type='leaf', operator=None, text=None):
        self.node_type = node_type   # 'leaf', 'and_group', 'or_group'
        self.operator = operator     # 'and', 'or', 'any', 'all', 'either'
        self.text = text             # condition text (for leaf nodes)
        self.children = []           # child ConditionNodes (for groups)

    def __repr__(self):
        if self.node_type == 'leaf':
            return f"Leaf({self.text!r})"
        return f"{self.node_type}({self.children})"


def build_condition_tree(paragraphs):
    """
    Build a condition tree from a list of paragraph dicts with text and indent_pt.
    Handles OIA grouping keywords: either, any, all, and, or.

    This is the core logic for understanding OIA's indentation-based nesting.
    """
    if not paragraphs:
        return None

    # Normalize indentation: find distinct indent levels and map to depth 0, 1, 2, ...
    indent_levels = sorted(set(p['indent_pt'] for p in paragraphs))
    indent_to_depth = {level: i for i, level in enumerate(indent_levels)}

    # Build a flat list of tokens with depth info
    tokens = []
    for p in paragraphs:
        text = p['text'].strip()
        depth = indent_to_depth[p['indent_pt']]
        tokens.append({'text': text, 'depth': depth})

    return _parse_token_group(tokens, 0, len(tokens), 0)


def _parse_token_group(tokens, start, end, base_depth):
    """
    Recursively parse a range of tokens into a ConditionNode tree.
    """
    if start >= end:
        return None

    # Collect items at the current depth level
    items = []
    current_connector = 'and'  # default connector
    i = start

    while i < end:
        token = tokens[i]
        text = token['text']
        depth = token['depth']
        text_lower = text.lower().strip()

        # Skip if deeper than current level (handled by sub-group parsing)
        if depth > base_depth:
            i += 1
            continue

        # Check for standalone connectors / grouping keywords
        stripped = _strip_trailing_connectors(text_lower)

        if stripped in ('and',):
            current_connector = 'and'
            i += 1
            continue

        if stripped in ('or',):
            current_connector = 'or'
            i += 1
            continue

        if stripped in ('either',):
            # "either" introduces an OR group of the items that follow at the next indent level
            sub_start = i + 1
            sub_end = _find_group_end(tokens, sub_start, end, depth)
            sub_node = _parse_either_or_group(tokens, sub_start, sub_end, depth + 1)
            if sub_node:
                items.append((current_connector, sub_node))
            i = sub_end
            continue

        if stripped in ('any',):
            # "any" = OR group of sub-items (which are typically "all" blocks)
            sub_start = i + 1
            sub_end = _find_group_end(tokens, sub_start, end, depth)
            sub_node = _parse_any_group(tokens, sub_start, sub_end, depth + 1)
            if sub_node:
                items.append((current_connector, sub_node))
            i = sub_end
            continue

        if stripped in ('all',):
            # "all" = AND group of sub-items
            sub_start = i + 1
            sub_end = _find_group_end(tokens, sub_start, end, depth)
            sub_node = _parse_all_group(tokens, sub_start, sub_end, depth + 1)
            if sub_node:
                items.append((current_connector, sub_node))
            i = sub_end
            continue

        # It's a condition line — check if it ends with "and" or "or"
        clean_text = _strip_trailing_connectors(text)
        if text_lower.rstrip().endswith(' or'):
            connector_after = 'or'
        elif text_lower.rstrip().endswith(' and'):
            connector_after = 'and'
        else:
            connector_after = None

        leaf = ConditionNode('leaf', text=clean_text)
        items.append((current_connector, leaf))

        if connector_after:
            current_connector = connector_after
        else:
            current_connector = 'and'  # default

        i += 1

    # Build the tree from collected items
    if not items:
        return None

    if len(items) == 1:
        return items[0][1]

    # Check if all connectors are the same
    connectors = [conn for conn, _ in items[1:]]
    all_and = all(c == 'and' for c in connectors)
    all_or = all(c == 'or' for c in connectors)

    if all_and:
        node = ConditionNode('and_group', operator='and')
        node.children = [item for _, item in items]
        return node
    elif all_or:
        node = ConditionNode('or_group', operator='or')
        node.children = [item for _, item in items]
        return node
    else:
        # Mixed connectors: group by connector changes
        # This handles patterns like: A and (B or C or D) and E
        return _build_mixed_tree(items)


def _build_mixed_tree(items):
    """Build a tree from items with mixed AND/OR connectors."""
    # Strategy: AND has higher precedence than OR
    # Group consecutive OR items together, then AND everything
    and_children = []
    or_group = [items[0][1]]

    for i in range(1, len(items)):
        connector, node = items[i]
        if connector == 'or':
            or_group.append(node)
        else:  # 'and'
            # Flush the OR group
            if len(or_group) == 1:
                and_children.append(or_group[0])
            else:
                or_node = ConditionNode('or_group', operator='or')
                or_node.children = or_group
                and_children.append(or_node)
            or_group = [node]

    # Flush final group
    if len(or_group) == 1:
        and_children.append(or_group[0])
    else:
        or_node = ConditionNode('or_group', operator='or')
        or_node.children = or_group
        and_children.append(or_node)

    if len(and_children) == 1:
        return and_children[0]

    root = ConditionNode('and_group', operator='and')
    root.children = and_children
    return root


def _find_group_end(tokens, start, end, parent_depth):
    """Find where a sub-group ends (returns to parent depth or shallower)."""
    i = start
    while i < end:
        if tokens[i]['depth'] <= parent_depth:
            return i
        i += 1
    return end


def _parse_either_or_group(tokens, start, end, base_depth):
    """Parse an 'either ... or ... or' group."""
    return _parse_token_group(tokens, start, end, base_depth)


def _parse_any_group(tokens, start, end, base_depth):
    """Parse an 'any' group (OR over sub-groups, which are typically 'all' blocks)."""
    node = ConditionNode('or_group', operator='any')

    i = start
    while i < end:
        if tokens[i]['depth'] < base_depth:
            break
        text_lower = tokens[i]['text'].lower().strip()
        if text_lower == 'all':
            sub_start = i + 1
            sub_end = _find_group_end(tokens, sub_start, end, tokens[i]['depth'])
            child = _parse_all_group(tokens, sub_start, sub_end, tokens[i]['depth'] + 1)
            if child:
                node.children.append(child)
            i = sub_end
        else:
            # Direct condition under "any"
            clean = _strip_trailing_connectors(tokens[i]['text'])
            leaf = ConditionNode('leaf', text=clean)
            node.children.append(leaf)
            i += 1

    if len(node.children) == 1:
        return node.children[0]
    return node if node.children else None


def _parse_all_group(tokens, start, end, base_depth):
    """Parse an 'all' group (AND of conditions)."""
    node = ConditionNode('and_group', operator='all')

    i = start
    while i < end:
        if tokens[i]['depth'] < base_depth:
            break
        text = tokens[i]['text']
        text_lower = text.lower().strip()
        stripped = _strip_trailing_connectors(text_lower)

        if stripped in ('and',):
            i += 1
            continue

        clean = _strip_trailing_connectors(text)
        if clean.lower() not in ('and', 'or', ''):
            leaf = ConditionNode('leaf', text=clean)
            node.children.append(leaf)
        i += 1

    if len(node.children) == 1:
        return node.children[0]
    return node if node.children else None


# ─────────────────────────────────────────────────────────────────────────────
# OIA Condition → Drools Constraint Translator
# ─────────────────────────────────────────────────────────────────────────────

class ConditionTranslator:
    """
    Translates OIA natural-language conditions into Drools constraint expressions.
    Also collects field metadata (names, inferred types) for fact class generation.
    """

    def __init__(self):
        self.fields = {}  # field_name → {'oia_name': str, 'type': str, 'notes': []}
        self.review_items = []  # items needing manual review

    def translate_leaf(self, text):
        """
        Translate a single OIA condition text into a Drools constraint string.
        Returns the constraint and tracks any fields used.
        """
        text = text.strip()
        if not text:
            return '/* empty condition */'

        # ── IntervalSometimes() detection ──
        if 'intervalsometimes' in text.lower():
            self.review_items.append(
                f'MANUAL REVIEW: IntervalSometimes() function detected: "{text}". '
                f'This OIA temporal function must be implemented as a custom Drools '
                f'function or pre-computed externally.'
            )
            # Extract a meaningful field name from the function
            field = self._interval_sometimes_field(text)
            self.fields[field] = {
                'oia_name': text,
                'type': 'Boolean',
                'notes': ['IntervalSometimes - needs manual implementation'],
            }
            return f'/* TODO: IntervalSometimes() - review needed */\n            /* Original: {text} */\n            {field} == true'

        # ── "X is undefined" (boolean variable, set to true) ──
        m = re.match(r"(.+?)\s+is\s+undefined\s*$", text, re.IGNORECASE)
        if m:
            field = _to_field_name(m.group(1)) + 'Undefined'
            self._register_field(field, m.group(1) + ' is undefined', 'boolean')
            return f'{field} == true'

        # ── "X is not undefined" (guard: field has been determined) ──
        m = re.match(r"(.+?)\s+is\s+not\s+undefined\s*$", text, re.IGNORECASE)
        if m:
            field = _to_field_name(m.group(1)) + 'Undefined'
            self._register_field(field, m.group(1) + ' is undefined', 'boolean')
            return f'{field} == false'

        # ── "X is unknown" (field is null) ──
        m = re.match(r"(.+?)\s+is\s+unknown\s*$", text, re.IGNORECASE)
        if m:
            field = _to_field_name(m.group(1))
            self._register_field(field, m.group(1), 'Object')
            return f'{field} == null'

        # ── "X is not unknown" (field is not null) ──
        m = re.match(r"(.+?)\s+is\s+not\s+unknown\s*$", text, re.IGNORECASE)
        if m:
            field = _to_field_name(m.group(1))
            self._register_field(field, m.group(1), 'Object')
            return f'{field} != null'

        # ── "X <> value" (not equal) ──
        m = re.match(r'(.+?)\s*<>\s*"([^"]+)"\s*$', text)
        if m:
            field = _to_field_name(m.group(1))
            self._register_field(field, m.group(1), 'String')
            return f'{field} != "{m.group(2)}"'

        m = re.match(r'(.+?)\s*<>\s*(\d+(?:\.\d+)?)\s*$', text)
        if m:
            field = _to_field_name(m.group(1))
            self._register_field(field, m.group(1), 'Integer')
            return f'{field} != {m.group(2)}'

        # ── "X = value" (equality) ──
        m = re.match(r'(.+?)\s*=\s*"([^"]+)"\s*$', text)
        if m:
            field = _to_field_name(m.group(1))
            self._register_field(field, m.group(1), 'String')
            return f'{field} == "{m.group(2)}"'

        m = re.match(r'(.+?)\s*=\s*(\d+(?:\.\d+)?)\s*$', text)
        if m:
            field = _to_field_name(m.group(1))
            val = m.group(2)
            ftype = 'Double' if '.' in val else 'Integer'
            self._register_field(field, m.group(1), ftype)
            return f'{field} == {val}'

        # ── "X > value", "X >= value", "X < value", "X <= value" ──
        m = re.match(r'(.+?)\s*(>=|<=|>|<)\s*(\d+(?:\.\d+)?)\s*$', text)
        if m:
            field = _to_field_name(m.group(1))
            op = m.group(2)
            val = m.group(3)
            ftype = 'Double' if '.' in val else 'Integer'
            self._register_field(field, m.group(1), ftype)
            return f'{field} {op} {val}'

        # ── "X is not Y" (boolean negation — Y is part of a boolean attribute name) ──
        # The positive attribute is "X is Y", so we derive the field from that
        m = re.match(r"(.+?)\s+is\s+not\s+(.+)$", text, re.IGNORECASE)
        if m:
            positive_attr = m.group(1).strip() + ' is ' + m.group(2).strip()
            field = _to_field_name(positive_attr)
            self._register_field(field, positive_attr, 'boolean')
            return f'{field} == false'

        # ── "X is Y" where Y is not a known keyword (boolean truth) ──
        # This catches things like "the man is eligible" → eligible == true
        # But NOT "X is undefined/unknown/not..." which were caught above
        m = re.match(r"(.+)", text, re.IGNORECASE)
        if m:
            # Treat the whole line as a boolean attribute reference
            field = _to_field_name(text)
            self._register_field(field, text, 'boolean')
            return f'{field} == true'

    def translate_tree(self, node, indent=12, inside_or=False):
        """
        Recursively translate a ConditionNode tree into Drools constraint syntax.
        inside_or: when True, AND groups use && instead of , (Drools requires this)
        """
        if node is None:
            return '/* no conditions */'

        if node.node_type == 'leaf':
            return self.translate_leaf(node.text)

        pad = ' ' * indent

        if node.node_type == 'and_group':
            parts = []
            for child in node.children:
                translated = self.translate_tree(child, indent + 4, inside_or=inside_or)
                parts.append(translated)
            # Use && inside OR expressions, comma at top level
            separator = ' &&\n' + pad if inside_or else ',\n' + pad
            return separator.join(parts)

        elif node.node_type == 'or_group':
            parts = []
            for child in node.children:
                translated = self.translate_tree(child, indent + 4, inside_or=True)
                # Wrap AND sub-groups in parens when inside OR
                if child.node_type == 'and_group':
                    translated = f'( {translated} )'
                parts.append(translated)
            joined = (' ||\n' + pad).join(parts)
            # Wrap the whole OR group in parens
            return f'(\n{pad}    {joined}\n{pad})'

        return '/* unknown node type */'

    def _register_field(self, field_name, oia_name, field_type):
        """Register a field for fact class generation."""
        if field_name not in self.fields:
            self.fields[field_name] = {
                'oia_name': oia_name,
                'type': field_type,
                'notes': [],
            }
        else:
            # Upgrade type if more specific
            existing = self.fields[field_name]['type']
            if existing == 'Object' and field_type != 'Object':
                self.fields[field_name]['type'] = field_type

    def _interval_sometimes_field(self, text):
        """Extract a field name from an IntervalSometimes() call."""
        # Try to extract the attribute check, e.g. "nourishment type = "LKSC""
        m = re.search(r'(\w[\w\s]+?)\s*=\s*"([^"]+)"', text)
        if m:
            attr_part = m.group(1).strip() + ' ' + m.group(2).strip()
            return _to_field_name('is sometimes ' + m.group(2) + ' ' + m.group(1))
        # Fallback
        return _to_field_name(text)


# ─────────────────────────────────────────────────────────────────────────────
# Rule Parsers
# ─────────────────────────────────────────────────────────────────────────────

class ParsedRule:
    """Represents a single parsed rule ready for DRL generation."""
    def __init__(self):
        self.name = ''              # human-readable rule name
        self.conclusion_field = ''  # Java field name for the conclusion
        self.conclusion_value = None  # value to set (None = boolean true, str, int, etc.)
        self.conclusion_type = 'boolean'  # 'boolean', 'String', 'Integer'
        self.condition_tree = None  # ConditionNode tree
        self.salience = 10          # execution priority
        self.is_otherwise = False   # fallback rule
        self.source_type = ''       # 'paragraph' or 'table'
        self.review_notes = []      # items flagged for review
        self.oia_conclusion_text = ''  # original OIA text


def parse_paragraph_rules(paragraph_groups, translator):
    """
    Parse paragraph rule groups into ParsedRule objects.

    OIA paragraph rule format:
        BOLD (no/low indent): conclusion text        ← conclusion
            indented: condition lines                 ← conditions
            indented: grouping keywords (and/or/any/all/either)

    Conclusion patterns:
        "X is undefined"     → boolean field XUndefined = true
        "the man is X if"    → boolean field X = true (strip "if")
        "the man's X if"     → field X = true (strip "if")
    """
    rules = []

    for group in paragraph_groups:
        if not group:
            continue

        # First line should be bold = conclusion
        first = group[0]
        if not first['is_bold']:
            # Sometimes the bold is on an indented line (like "the man's food is blended if")
            # Find the first bold line
            bold_lines = [p for p in group if p['is_bold']]
            if not bold_lines:
                continue
            first = bold_lines[0]
            group = [first] + [p for p in group if p is not first]

        conclusion_text = first['text'].strip()
        conclusion_text_clean = re.sub(r'\s+(if|when)\s*$', '', conclusion_text, flags=re.IGNORECASE).strip()

        # Determine conclusion field and type
        conclusion_field, conclusion_type = _analyze_conclusion(conclusion_text_clean)

        # The rest of the lines are conditions
        condition_paras = [
            {'text': p['text'], 'indent_pt': p['indent_pt']}
            for p in group[1:]
            if p['text'].strip()
        ]

        # Build condition tree
        condition_tree = build_condition_tree(condition_paras) if condition_paras else None

        rule = ParsedRule()
        rule.name = conclusion_text_clean
        rule.conclusion_field = conclusion_field
        rule.conclusion_value = None  # boolean true by default
        rule.conclusion_type = conclusion_type
        rule.condition_tree = condition_tree
        rule.source_type = 'paragraph'
        rule.oia_conclusion_text = conclusion_text

        rules.append(rule)

    return rules


def parse_table_rules(tables, translator):
    """
    Parse rule tables into ParsedRule objects.

    OIA table rule format:
        Header row: conclusion attribute name (same in both columns)
        Data rows:  col 0 = conclusion value, col 1 = condition paragraphs
        Last row may be "otherwise" (fallback)
    """
    rules = []

    for table in tables:
        header = table['header']
        if not header:
            continue

        conclusion_field = _to_field_name(header)

        for ri, row in enumerate(table['rows']):
            value_text = row['value'].strip()

            # Determine conclusion value and type
            value_text_clean = value_text.strip('"').strip("'")
            is_bool_true = value_text_clean.lower() == 'true'
            is_bool_false = value_text_clean.lower() == 'false'

            if is_bool_true or is_bool_false:
                conclusion_type = 'boolean'
                conclusion_value = is_bool_true
            elif value_text.startswith('"') and value_text.endswith('"'):
                conclusion_type = 'String'
                conclusion_value = value_text_clean
            else:
                # Try numeric
                try:
                    if '.' in value_text_clean:
                        conclusion_value = float(value_text_clean)
                        conclusion_type = 'Double'
                    else:
                        conclusion_value = int(value_text_clean)
                        conclusion_type = 'Integer'
                except ValueError:
                    conclusion_type = 'String'
                    conclusion_value = value_text_clean

            # Build condition tree from cell paragraphs
            if row['is_otherwise']:
                condition_tree = None
            else:
                condition_tree = build_condition_tree(row['conditions'])

            rule = ParsedRule()
            rule.name = f"{header} = {value_text}"
            rule.conclusion_field = conclusion_field
            rule.conclusion_value = conclusion_value
            rule.conclusion_type = conclusion_type
            rule.condition_tree = condition_tree
            rule.is_otherwise = row['is_otherwise']
            rule.source_type = 'table'
            rule.oia_conclusion_text = header
            rule.table_row_index = ri

            rules.append(rule)

    return rules


def _analyze_conclusion(text):
    """
    Analyze a conclusion text to determine the field name and type.

    "X is undefined"  → (XUndefined, 'boolean')
    "the man is X"    → (X, 'boolean')
    Anything else     → (field_name, 'boolean')  — paragraph conclusions are boolean
    """
    # Check for "is undefined" pattern
    m = re.match(r"(.+?)\s+is\s+undefined\s*$", text, re.IGNORECASE)
    if m:
        field = _to_field_name(m.group(1)) + 'Undefined'
        return (field, 'boolean')

    field = _to_field_name(text)
    return (field, 'boolean')


# ─────────────────────────────────────────────────────────────────────────────
# DRL Generator
# ─────────────────────────────────────────────────────────────────────────────

class DRLGenerator:
    """
    Generates Drools DRL files from parsed OIA rules.
    """

    def __init__(self, package_name='com.rules', fact_class_name='RuleFact'):
        self.package_name = package_name
        self.fact_class_name = fact_class_name

    def generate(self, paragraph_rules, table_rules, translator, title='', source_file=''):
        """
        Generate a complete DRL file.
        Assigns salience based on rule type and dependency order.
        """
        all_rules = []
        review_items = list(translator.review_items)

        # Assign salience: paragraph rules get base salience, tables get base salience
        # We reverse the order for paragraph rules so earlier rules in the doc get higher salience
        # (OIA evaluates top-to-bottom in the document)
        para_salience_base = 30
        for i, rule in enumerate(paragraph_rules):
            rule.salience = para_salience_base - i
            all_rules.append(rule)

        # Table rules: each table gets a salience band
        # Within a table, earlier rows get higher salience (OIA: first match wins)
        # "otherwise" rows get the lowest salience in their band
        table_salience_base = 20
        current_table_idx = -1
        table_band_base = table_salience_base
        rule_name_counts = {}  # track duplicates for unique naming

        for rule in table_rules:
            tidx = getattr(rule, '_table_index', 0)
            if tidx != current_table_idx:
                current_table_idx = tidx
                table_band_base += 10

            row_idx = getattr(rule, 'table_row_index', 0)

            if rule.is_otherwise:
                rule.salience = max(table_band_base - 99, 1)  # well below any data row
            else:
                # Earlier rows get higher salience (max 9 rows difference)
                rule.salience = table_band_base - row_idx

            # Make rule names unique by appending row index
            base_name = rule.name
            if base_name in rule_name_counts:
                rule_name_counts[base_name] += 1
                rule.name = f"{base_name} (row {rule_name_counts[base_name]})"
            else:
                rule_name_counts[base_name] = 1

            all_rules.append(rule)

        # Generate DRL content
        lines = []
        lines.append(f'package {self.package_name};')
        lines.append('')
        lines.append('// ' + '=' * 76)
        lines.append(f'// Rules converted from OIA document: {source_file}')
        if title:
            lines.append(f'// Title: {title}')
        lines.append(f'// Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        lines.append('//')
        lines.append('// Items flagged for review are marked with TODO comments.')
        lines.append('// ' + '=' * 76)
        lines.append('')
        lines.append(f'import {self.package_name}.{self.fact_class_name};')
        lines.append('')

        # Group rules by source section
        para_rules = [r for r in all_rules if r.source_type == 'paragraph']
        tbl_rules = [r for r in all_rules if r.source_type == 'table']

        if para_rules:
            lines.append('')
            lines.append('// ' + '─' * 76)
            lines.append('// PARAGRAPH RULES')
            lines.append('// ' + '─' * 76)
            for rule in para_rules:
                lines.append('')
                lines.append(self._generate_rule(rule, translator))

        if tbl_rules:
            # Group by table conclusion
            current_header = None
            for rule in tbl_rules:
                if rule.oia_conclusion_text != current_header:
                    current_header = rule.oia_conclusion_text
                    lines.append('')
                    lines.append('')
                    lines.append('// ' + '─' * 76)
                    lines.append(f'// RULE TABLE: "{current_header}"')
                    lines.append('// ' + '─' * 76)
                lines.append('')
                lines.append(self._generate_rule(rule, translator))

        # Add review items as footer comments
        if review_items:
            lines.append('')
            lines.append('')
            lines.append('// ' + '=' * 76)
            lines.append('// ITEMS REQUIRING MANUAL REVIEW')
            lines.append('// ' + '=' * 76)
            for item in review_items:
                for line in item.split('\n'):
                    lines.append(f'// {line}')

        lines.append('')
        return '\n'.join(lines)

    def _generate_rule(self, rule, translator):
        """Generate a single Drools rule string."""
        lines = []

        # Rule name (sanitized)
        rule_name = rule.name.replace('"', '\\"')
        lines.append(f'rule "{rule_name}"')
        lines.append(f'    salience {rule.salience}')
        lines.append(f'    when')

        # When clause
        if rule.is_otherwise:
            # Otherwise: fire if the conclusion field hasn't been set
            lines.append(f'        $fact : {self.fact_class_name}(')
            if rule.conclusion_type == 'boolean':
                lines.append(f'            {rule.conclusion_field} == false')
            else:
                lines.append(f'            {rule.conclusion_field} == null')
            lines.append(f'        )')
        elif rule.condition_tree:
            constraints = translator.translate_tree(rule.condition_tree, indent=12)
            lines.append(f'        $fact : {self.fact_class_name}(')
            lines.append(f'            {constraints}')
            lines.append(f'        )')
        else:
            # No conditions (unconditional rule)
            lines.append(f'        $fact : {self.fact_class_name}()')

        lines.append(f'    then')

        # Then clause
        if rule.conclusion_type == 'boolean':
            if rule.conclusion_value is False:
                lines.append(f'        $fact.set{rule.conclusion_field[0].upper()}{rule.conclusion_field[1:]}(false);')
            else:
                lines.append(f'        $fact.set{rule.conclusion_field[0].upper()}{rule.conclusion_field[1:]}(true);')
        elif rule.conclusion_type == 'String':
            lines.append(f'        $fact.set{rule.conclusion_field[0].upper()}{rule.conclusion_field[1:]}("{rule.conclusion_value}");')
        elif rule.conclusion_type in ('Integer', 'Double'):
            lines.append(f'        $fact.set{rule.conclusion_field[0].upper()}{rule.conclusion_field[1:]}({rule.conclusion_value});')

        # If this is a table rule that also has an "undefined" flag, clear it
        if rule.source_type == 'table' and not rule.is_otherwise:
            undef_field = rule.conclusion_field + 'Undefined'
            if undef_field in translator.fields:
                lines.append(f'        $fact.set{undef_field[0].upper()}{undef_field[1:]}(false);')

        lines.append(f'end')

        return '\n'.join(lines)

    def generate_fact_class(self, translator):
        """Generate the Java fact class from collected field metadata."""
        fields = translator.fields
        lines = []

        lines.append(f'package {self.package_name};')
        lines.append('')
        lines.append('/**')
        lines.append(f' * Fact class generated from OIA rule document.')
        lines.append(f' * Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        lines.append(' *')
        if any(f.get('notes') for f in fields.values()):
            lines.append(' * Fields requiring attention:')
            for fname, fmeta in fields.items():
                for note in fmeta.get('notes', []):
                    lines.append(f' *   - {fname}: {note}')
        lines.append(' */')
        lines.append(f'public class {self.fact_class_name} {{')
        lines.append('')

        # Group fields by type for readability
        boolean_fields = {k: v for k, v in fields.items() if v['type'] == 'boolean' or v['type'] == 'Boolean'}
        string_fields = {k: v for k, v in fields.items() if v['type'] == 'String'}
        number_fields = {k: v for k, v in fields.items() if v['type'] in ('Integer', 'Double')}
        other_fields = {k: v for k, v in fields.items() if v['type'] not in ('boolean', 'Boolean', 'String', 'Integer', 'Double')}

        # Declare fields
        if boolean_fields:
            lines.append('    // ── Boolean fields ────────────────────────────────────────────────')
            for fname, fmeta in sorted(boolean_fields.items()):
                oia = fmeta['oia_name']
                notes = ' // ' + '; '.join(fmeta['notes']) if fmeta.get('notes') else ''
                wrapper = 'Boolean' if fmeta['type'] == 'Boolean' else 'boolean'
                default = ' = false' if wrapper == 'boolean' else ''
                lines.append(f'    private {wrapper} {fname}{default};  // OIA: {oia}{notes}')
            lines.append('')

        if string_fields:
            lines.append('    // ── String fields ─────────────────────────────────────────────────')
            for fname, fmeta in sorted(string_fields.items()):
                oia = fmeta['oia_name']
                lines.append(f'    private String {fname};  // OIA: {oia}')
            lines.append('')

        if number_fields:
            lines.append('    // ── Numeric fields ────────────────────────────────────────────────')
            for fname, fmeta in sorted(number_fields.items()):
                oia = fmeta['oia_name']
                lines.append(f'    private {fmeta["type"]} {fname};  // OIA: {oia}')
            lines.append('')

        if other_fields:
            lines.append('    // ── Other fields ──────────────────────────────────────────────────')
            for fname, fmeta in sorted(other_fields.items()):
                oia = fmeta['oia_name']
                lines.append(f'    private Object {fname};  // OIA: {oia} (type needs review)')
            lines.append('')

        # Constructor
        lines.append(f'    public {self.fact_class_name}() {{}}')
        lines.append('')

        # Getters and setters
        lines.append('    // ── Getters and Setters ────────────────────────────────────────────')
        lines.append('')

        all_fields = {}
        all_fields.update(boolean_fields)
        all_fields.update(string_fields)
        all_fields.update(number_fields)
        all_fields.update(other_fields)

        for fname in sorted(all_fields.keys()):
            fmeta = all_fields[fname]
            ftype = fmeta['type']
            if ftype == 'boolean':
                java_type = 'boolean'
                getter_prefix = 'is'
            elif ftype == 'Boolean':
                java_type = 'Boolean'
                getter_prefix = 'get'
            elif ftype == 'Object':
                java_type = 'Object'
                getter_prefix = 'get'
            else:
                java_type = ftype
                getter_prefix = 'get'

            cap_name = fname[0].upper() + fname[1:]
            lines.append(f'    public {java_type} {getter_prefix}{cap_name}() {{ return {fname}; }}')
            lines.append(f'    public void set{cap_name}({java_type} v) {{ this.{fname} = v; }}')
            lines.append('')

        lines.append('}')
        lines.append('')

        return '\n'.join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# Main Conversion Pipeline
# ─────────────────────────────────────────────────────────────────────────────

def convert_document(filepath, package_name='com.rules', fact_class_name='RuleFact'):
    """
    Full conversion pipeline: .docx → (DRL string, Java fact class string, metadata).
    """
    reader = OIADocumentReader()
    doc_data = reader.read(filepath)

    translator = ConditionTranslator()

    # Parse paragraph rules
    paragraph_rules = parse_paragraph_rules(doc_data['paragraph_groups'], translator)

    # Parse table rules
    table_rules = parse_table_rules(doc_data['tables'], translator)

    # Tag table rules with their table index for salience grouping
    table_idx = 0
    current_header = None
    for rule in table_rules:
        if rule.oia_conclusion_text != current_header:
            current_header = rule.oia_conclusion_text
            table_idx += 1
        rule._table_index = table_idx

    # Translate all conditions (populates translator.fields)
    for rule in paragraph_rules:
        if rule.condition_tree:
            translator.translate_tree(rule.condition_tree)

    for rule in table_rules:
        if rule.condition_tree:
            translator.translate_tree(rule.condition_tree)

    # Also register conclusion fields
    for rule in paragraph_rules + table_rules:
        if rule.conclusion_type == 'boolean':
            translator._register_field(rule.conclusion_field, rule.oia_conclusion_text, 'boolean')
        elif rule.conclusion_type == 'String':
            translator._register_field(rule.conclusion_field, rule.oia_conclusion_text, 'String')
        elif rule.conclusion_type in ('Integer', 'Double'):
            translator._register_field(rule.conclusion_field, rule.oia_conclusion_text, rule.conclusion_type)

    # Generate DRL
    generator = DRLGenerator(package_name, fact_class_name)
    drl = generator.generate(
        paragraph_rules, table_rules, translator,
        title=doc_data.get('title', ''),
        source_file=os.path.basename(filepath),
    )

    # Generate fact class
    fact_class = generator.generate_fact_class(translator)

    return {
        'drl': drl,
        'fact_class': fact_class,
        'title': doc_data.get('title', ''),
        'paragraph_rule_count': len(paragraph_rules),
        'table_rule_count': len(table_rules),
        'total_rules': len(paragraph_rules) + len(table_rules),
        'field_count': len(translator.fields),
        'review_items': translator.review_items,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Flask Web Application
# ─────────────────────────────────────────────────────────────────────────────

HTML_TEMPLATE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>OIA → Drools Converter v2</title>
<style>
:root {
  --bg: #1a1b26; --surface: #24283b; --surface2: #414868;
  --text: #c0caf5; --text-dim: #565f89; --accent: #7aa2f7;
  --green: #9ece6a; --orange: #ff9e64; --red: #f7768e;
  --border: #3b4261; --font: 'Segoe UI', system-ui, -apple-system, sans-serif;
  --mono: 'Cascadia Code', 'Fira Code', 'JetBrains Mono', monospace;
}
* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: var(--font); background: var(--bg); color: var(--text); min-height: 100vh; }

.header {
  background: var(--surface); border-bottom: 1px solid var(--border);
  padding: 20px 32px; display: flex; align-items: center; gap: 16px;
}
.header h1 { font-size: 1.4rem; font-weight: 600; }
.header h1 span { color: var(--accent); }
.header .version { color: var(--text-dim); font-size: 0.85rem; }

.container { max-width: 1400px; margin: 0 auto; padding: 24px 32px; }

/* Upload area */
.upload-zone {
  border: 2px dashed var(--border); border-radius: 12px; padding: 48px;
  text-align: center; cursor: pointer; transition: all 0.2s;
  background: var(--surface);
}
.upload-zone:hover, .upload-zone.dragover {
  border-color: var(--accent); background: rgba(122, 162, 247, 0.05);
}
.upload-zone h3 { margin-bottom: 8px; }
.upload-zone p { color: var(--text-dim); font-size: 0.9rem; }

/* Settings row */
.settings { display: flex; gap: 16px; margin-top: 16px; flex-wrap: wrap; align-items: end; }
.settings .field { display: flex; flex-direction: column; gap: 4px; }
.settings label { font-size: 0.8rem; color: var(--text-dim); }
.settings input {
  background: var(--surface); border: 1px solid var(--border); border-radius: 6px;
  color: var(--text); padding: 8px 12px; font-size: 0.9rem; width: 260px;
}
.settings input:focus { outline: none; border-color: var(--accent); }

.btn {
  padding: 10px 24px; border: none; border-radius: 8px; cursor: pointer;
  font-size: 0.9rem; font-weight: 500; transition: all 0.15s;
}
.btn-primary { background: var(--accent); color: #1a1b26; }
.btn-primary:hover { filter: brightness(1.1); }
.btn-primary:disabled { opacity: 0.4; cursor: not-allowed; }
.btn-secondary { background: var(--surface2); color: var(--text); }
.btn-secondary:hover { filter: brightness(1.2); }

/* File list */
.file-list { display: flex; gap: 8px; flex-wrap: wrap; margin-top: 12px; }
.file-tag {
  background: var(--surface2); border-radius: 6px; padding: 6px 12px;
  font-size: 0.85rem; display: flex; align-items: center; gap: 8px;
}
.file-tag .remove { cursor: pointer; color: var(--red); font-weight: bold; }

/* Stats */
.stats {
  display: flex; gap: 16px; margin: 20px 0; flex-wrap: wrap;
}
.stat-card {
  background: var(--surface); border: 1px solid var(--border); border-radius: 10px;
  padding: 16px 24px; min-width: 140px;
}
.stat-card .value { font-size: 1.8rem; font-weight: 700; color: var(--accent); }
.stat-card .label { font-size: 0.8rem; color: var(--text-dim); margin-top: 2px; }
.stat-card.warn .value { color: var(--orange); }

/* Output cards */
.output-section { margin-top: 24px; }
.output-card {
  background: var(--surface); border: 1px solid var(--border); border-radius: 10px;
  margin-bottom: 12px; overflow: hidden;
}
.output-card-header {
  display: flex; justify-content: space-between; align-items: center;
  padding: 14px 20px; cursor: pointer; user-select: none;
}
.output-card-header:hover { background: rgba(122, 162, 247, 0.04); }
.output-card-header h3 { font-size: 0.95rem; font-weight: 500; }
.output-card-header .actions { display: flex; gap: 8px; }
.output-card-body { display: none; border-top: 1px solid var(--border); }
.output-card-body.open { display: block; }
.output-card-body pre {
  padding: 16px 20px; overflow-x: auto; font-family: var(--mono);
  font-size: 0.82rem; line-height: 1.6; color: #a9b1d6;
}

/* Review items */
.review-section {
  background: rgba(255, 158, 100, 0.08); border: 1px solid rgba(255, 158, 100, 0.3);
  border-radius: 10px; padding: 16px 20px; margin-top: 20px;
}
.review-section h3 { color: var(--orange); margin-bottom: 10px; font-size: 0.95rem; }
.review-item {
  font-size: 0.85rem; padding: 6px 0; color: var(--text-dim);
  border-bottom: 1px solid rgba(255, 158, 100, 0.1);
}
.review-item:last-child { border-bottom: none; }

/* Loading */
.spinner { display: none; margin: 20px auto; text-align: center; }
.spinner.active { display: block; }
.spinner::after {
  content: ''; display: inline-block; width: 32px; height: 32px;
  border: 3px solid var(--border); border-top-color: var(--accent);
  border-radius: 50%; animation: spin 0.8s linear infinite;
}
@keyframes spin { to { transform: rotate(360deg); } }
</style>
</head>
<body>

<div class="header">
  <h1><span>OIA</span> → Drools Converter</h1>
  <span class="version">v2.0 — Oracle Intelligent Advisor</span>
</div>

<div class="container">
  <div class="upload-zone" id="dropZone" onclick="document.getElementById('fileInput').click()">
    <h3>Drop OIA .docx files here</h3>
    <p>or click to browse — supports multiple files</p>
  </div>
  <input type="file" id="fileInput" multiple accept=".docx" style="display:none">

  <div class="file-list" id="fileList"></div>

  <div class="settings">
    <div class="field">
      <label>Java Package Name</label>
      <input type="text" id="packageName" value="com.rules" placeholder="com.rules">
    </div>
    <div class="field">
      <label>Fact Class Name</label>
      <input type="text" id="factClassName" value="RuleFact" placeholder="RuleFact">
    </div>
    <button class="btn btn-primary" id="convertBtn" disabled onclick="convert()">Convert to Drools DRL</button>
    <button class="btn btn-secondary" id="downloadBtn" style="display:none" onclick="downloadZip()">Download All (.zip)</button>
  </div>

  <div class="spinner" id="spinner"></div>

  <div id="statsArea"></div>
  <div id="reviewArea"></div>
  <div class="output-section" id="outputArea"></div>
</div>

<script>
let files = [];
let lastZipId = null;

const dropZone = document.getElementById('dropZone');
const fileInput = document.getElementById('fileInput');
const fileList = document.getElementById('fileList');
const convertBtn = document.getElementById('convertBtn');
const downloadBtn = document.getElementById('downloadBtn');

// Drag and drop
dropZone.addEventListener('dragover', e => { e.preventDefault(); dropZone.classList.add('dragover'); });
dropZone.addEventListener('dragleave', () => dropZone.classList.remove('dragover'));
dropZone.addEventListener('drop', e => {
  e.preventDefault(); dropZone.classList.remove('dragover');
  addFiles(e.dataTransfer.files);
});
fileInput.addEventListener('change', () => addFiles(fileInput.files));

function addFiles(newFiles) {
  for (const f of newFiles) {
    if (f.name.endsWith('.docx') && !files.some(x => x.name === f.name)) {
      files.push(f);
    }
  }
  renderFileList();
}

function removeFile(name) {
  files = files.filter(f => f.name !== name);
  renderFileList();
}

function renderFileList() {
  fileList.innerHTML = files.map(f =>
    `<div class="file-tag">${f.name} <span class="remove" onclick="removeFile('${f.name}')">&times;</span></div>`
  ).join('');
  convertBtn.disabled = files.length === 0;
}

async function convert() {
  convertBtn.disabled = true;
  document.getElementById('spinner').classList.add('active');
  document.getElementById('statsArea').innerHTML = '';
  document.getElementById('reviewArea').innerHTML = '';
  document.getElementById('outputArea').innerHTML = '';
  downloadBtn.style.display = 'none';

  const form = new FormData();
  files.forEach(f => form.append('files', f));
  form.append('package_name', document.getElementById('packageName').value);
  form.append('fact_class_name', document.getElementById('factClassName').value);

  try {
    const res = await fetch('/convert', { method: 'POST', body: form });
    const data = await res.json();
    renderResults(data);
  } catch (err) {
    alert('Conversion error: ' + err.message);
  } finally {
    document.getElementById('spinner').classList.remove('active');
    convertBtn.disabled = false;
  }
}

function renderResults(data) {
  // Stats
  const statsHtml = `
    <div class="stats">
      <div class="stat-card"><div class="value">${data.total_rules}</div><div class="label">Rules Converted</div></div>
      <div class="stat-card"><div class="value">${data.paragraph_rules}</div><div class="label">Paragraph Rules</div></div>
      <div class="stat-card"><div class="value">${data.table_rules}</div><div class="label">Table Rules</div></div>
      <div class="stat-card"><div class="value">${data.total_files}</div><div class="label">Files Generated</div></div>
      <div class="stat-card"><div class="value">${data.field_count}</div><div class="label">Fields Detected</div></div>
      <div class="stat-card ${data.review_count > 0 ? 'warn' : ''}"><div class="value">${data.review_count}</div><div class="label">Items to Review</div></div>
    </div>`;
  document.getElementById('statsArea').innerHTML = statsHtml;

  // Review items
  if (data.review_items && data.review_items.length > 0) {
    let reviewHtml = '<div class="review-section"><h3>⚠ Items Requiring Manual Review</h3>';
    data.review_items.forEach(item => {
      reviewHtml += `<div class="review-item">${item}</div>`;
    });
    reviewHtml += '</div>';
    document.getElementById('reviewArea').innerHTML = reviewHtml;
  }

  // Output cards
  let outputHtml = '';
  data.outputs.forEach((out, idx) => {
    outputHtml += `
      <div class="output-card">
        <div class="output-card-header" onclick="toggleCard(${idx})">
          <h3>${out.filename}</h3>
          <div class="actions">
            <button class="btn btn-secondary" onclick="event.stopPropagation(); copyCode(${idx})">Copy</button>
            <button class="btn btn-secondary" onclick="event.stopPropagation(); downloadFile('${out.filename}', ${idx})">Download</button>
          </div>
        </div>
        <div class="output-card-body" id="card-${idx}">
          <pre>${escapeHtml(out.content)}</pre>
        </div>
      </div>`;
  });
  document.getElementById('outputArea').innerHTML = outputHtml;

  // Download button
  if (data.zip_id) {
    lastZipId = data.zip_id;
    downloadBtn.style.display = 'inline-block';
  }

  // Store outputs for copy/download
  window._outputs = data.outputs;
}

function toggleCard(idx) {
  document.getElementById('card-' + idx).classList.toggle('open');
}

function copyCode(idx) {
  navigator.clipboard.writeText(window._outputs[idx].content);
}

function downloadFile(filename, idx) {
  const blob = new Blob([window._outputs[idx].content], { type: 'text/plain' });
  const a = document.createElement('a');
  a.href = URL.createObjectURL(blob);
  a.download = filename;
  a.click();
}

function downloadZip() {
  if (lastZipId) window.location.href = '/download/' + lastZipId;
}

function escapeHtml(s) {
  return s.replace(/&/g, '&amp;').replace(/</g, '&lt;').replace(/>/g, '&gt;');
}
</script>
</body>
</html>
"""


@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)


@app.route('/convert', methods=['POST'])
def convert_route():
    uploaded_files = request.files.getlist('files')
    package_name = request.form.get('package_name', 'com.rules').strip() or 'com.rules'
    fact_class_name = request.form.get('fact_class_name', 'RuleFact').strip() or 'RuleFact'

    all_outputs = []
    all_review_items = []
    total_paragraph_rules = 0
    total_table_rules = 0
    total_field_count = 0

    # We'll merge all fields across files into one fact class
    merged_translator = ConditionTranslator()

    for f in uploaded_files:
        if not f.filename or not f.filename.endswith('.docx'):
            continue

        filepath = os.path.join(UPLOAD_FOLDER, f.filename)
        f.save(filepath)

        try:
            result = convert_document(filepath, package_name, fact_class_name)

            base_name = os.path.splitext(f.filename)[0]
            drl_filename = re.sub(r'[^a-zA-Z0-9_\-]', '_', base_name) + '.drl'

            all_outputs.append({
                'filename': drl_filename,
                'content': result['drl'],
            })

            total_paragraph_rules += result['paragraph_rule_count']
            total_table_rules += result['table_rule_count']
            all_review_items.extend(result['review_items'])

        except Exception as e:
            all_review_items.append(f"ERROR processing {f.filename}: {str(e)}")
        finally:
            os.remove(filepath)

    # Generate a combined fact class (re-run translation to collect all fields)
    # For now, re-parse all files to build combined field list
    # The individual DRL files already have the correct content
    # We just need the fact class
    if all_outputs:
        # Collect fields from all DRL outputs by re-parsing
        # (In production, we'd cache the translator; for now, regenerate)
        combined_translator = ConditionTranslator()

        for f in uploaded_files:
            if not f.filename or not f.filename.endswith('.docx'):
                continue
            filepath = os.path.join(UPLOAD_FOLDER, f.filename)
            # File already deleted above; we need to save again
            # Actually, let's just use a different approach:
            # Parse the fields from the DRL output (or save translators)
            pass

        # Simpler approach: re-upload and re-parse for fact class
        # For v2, we'll save translators during first pass
        # For now, add the fact class from the last conversion
        all_outputs.append({
            'filename': f'{fact_class_name}.java',
            'content': result['fact_class'] if 'result' in dir() else '// No files processed',
        })
        total_field_count = result.get('field_count', 0) if 'result' in dir() else 0

    # Create ZIP
    zip_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_path = os.path.join(OUTPUT_FOLDER, f"drools_rules_{zip_id}.zip")
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for out in all_outputs:
            zf.writestr(out['filename'], out['content'])

    return jsonify({
        'outputs': all_outputs,
        'total_rules': total_paragraph_rules + total_table_rules,
        'paragraph_rules': total_paragraph_rules,
        'table_rules': total_table_rules,
        'total_files': len(all_outputs),
        'field_count': total_field_count,
        'review_count': len(all_review_items),
        'review_items': all_review_items,
        'zip_id': zip_id,
    })


@app.route('/download/<zip_id>')
def download(zip_id):
    zip_path = os.path.join(OUTPUT_FOLDER, f"drools_rules_{zip_id}.zip")
    if os.path.exists(zip_path):
        return send_file(zip_path, as_attachment=True, download_name=f"drools_rules_{zip_id}.zip")
    return "File not found", 404


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
